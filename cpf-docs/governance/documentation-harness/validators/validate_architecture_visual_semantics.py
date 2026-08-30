#!/usr/bin/env python3
import hashlib,json,sys
from pathlib import Path
H=Path(__file__).resolve().parents[1]
ROOT=H.parents[2]
CFG=H/'architecture-visual-semantics.json'

def fail(msg):
    print('ARCHITECTURE_VISUAL_SEMANTIC=FAIL '+msg)
    raise SystemExit(1)

def main():
    if not CFG.is_file(): fail('missing architecture-visual-semantics.json')
    d=json.loads(CFG.read_text(encoding='utf-8'))
    if d.get('harnessVersion')!='2.15.4': fail('harnessVersion mismatch')
    asset=ROOT/d.get('asset','')
    if not asset.is_file(): fail('asset missing '+str(asset))
    actual=hashlib.sha256(asset.read_bytes()).hexdigest().upper()
    if actual!=str(d.get('assetSha256','')).upper(): fail('asset SHA mismatch')
    nodes=d.get('nodes',{}); zones=d.get('zones',{})
    expected={
      'cpf-backoffice':'BUSINESS_DOMAIN_CPF',
      'cpf-backoffice-web':'CHANNEL_EDGE',
      'cpf-gateway':'CHANNEL_EDGE',
      'cpf-admin':'PLATFORM_OPERATIONS_RUNTIME',
      'cpf-batch':'PLATFORM_OPERATIONS_RUNTIME',
      'generated-domain':'BUSINESS_DOMAIN_CPF'
    }
    for node,zone in expected.items():
        if nodes.get(node,{}).get('zone')!=zone: fail(f'{node} must be in {zone}')
    for zid,z in zones.items():
        for node in z.get('forbidden',[]):
            if nodes.get(node,{}).get('zone')==zid: fail(f'forbidden node {node} placed in {zid}')
    hard=d.get('hardRules',{})
    for k in ['cpfBackofficeIsPlatformControlPlane','cpfBackofficeWebHasDatabase','cpfBackofficeWebDependsOnBusinessDomainJavaProject','internalDomainCallThroughGateway','backofficeAndBffConflated']:
        if hard.get(k) is not False: fail('hard rule must be false: '+k)
    flow=d.get('flows',{}).get('backofficeBrowserFlow',[])
    if flow!=['cpf-backoffice-web','cpf-gateway','cpf-backoffice']:
        fail('backofficeBrowserFlow must be cpf-backoffice-web -> cpf-gateway -> cpf-backoffice')
    # Cross-check current README companion explanation, not only sidecar metadata.
    readme=(ROOT/'README.md').read_text(encoding='utf-8')
    required=['cpf-backoffice-web','Optional Prebuilt Business Domain','플랫폼 Control Plane']
    for token in required:
        if token not in readme: fail('README architecture companion missing '+token)
    # Cross-check all official architecture-bearing DOCX files so figure semantics and formal deliverables cannot diverge.
    try:
        from docx import Document
        docs=[
            ROOT/'cpf-docs/deliverables/아키텍처설계서.docx',
            ROOT/'cpf-docs/guides/07_Specification_기술_명세.docx',
            ROOT/'cpf-docs/deliverables/기술사양서.docx',
        ]
        texts={}
        for dp in docs:
            doc=Document(dp)
            parts=[x.text for x in doc.paragraphs]
            for tab in doc.tables:
                for row in tab.rows:
                    parts.append(' | '.join(c.text for c in row.cells))
            texts[dp.name]='\n'.join(parts)
        arch=texts['아키텍처설계서.docx']
        for token in ['cpf-backoffice-web','Optional Prebuilt Business Domain','플랫폼 Control Plane']:
            if token not in arch: fail('Architecture Design missing '+token)
        spec=texts['07_Specification_기술_명세.docx']
        for token in ['Backoffice Web/BFF','cpf-backoffice-web','cpf-backoffice','플랫폼 Control Plane 아님']:
            if token not in spec: fail('Specification missing canonical Backoffice split: '+token)
        tech=texts['기술사양서.docx']
        for token in ['cpf-backoffice','Optional Prebuilt Business Domain','플랫폼 Control Plane 아님']:
            if token not in tech: fail('Technical Specification missing canonical Backoffice owner semantics: '+token)
        forbidden=['cpf-backoffice | 선택 Runtime/BFF','cpf-backoffice | Runtime/BFF','cpf-backoffice | Platform Control Plane']
        joined='\n'.join(texts.values())
        for phrase in forbidden:
            if phrase in joined: fail('conflated/misclassified Backoffice wording remains: '+phrase)
    except Exception as e:
        fail('official architecture document read failed: '+str(e))
    print('ARCHITECTURE_VISUAL_SEMANTIC=PASS')
    print('ASSET_SHA256='+actual)
    print('BACKOFFICE_ZONE=BUSINESS_DOMAIN_CPF')
    return 0
if __name__=='__main__': raise SystemExit(main())
