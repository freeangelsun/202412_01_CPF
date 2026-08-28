#!/usr/bin/env python3
import re,sys,zipfile,statistics,math
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
H=Path(__file__).resolve().parents[1]
ROOT=H.parents[2]
DOCS=[
 'cpf-docs/guides/02_프레임워크_개발자_가이드.docx','cpf-docs/guides/03_배치_개발자_가이드.docx',
 'cpf-docs/guides/04_운영자_매뉴얼.docx','cpf-docs/guides/05_배치_운영_가이드.docx',
 'cpf-docs/guides/06_Gateway_개발_사용_가이드.docx','cpf-docs/guides/07_Specification_기술_명세.docx',
 'cpf-docs/deliverables/아키텍처설계서.docx','cpf-docs/deliverables/기술사양서.docx',
 'cpf-docs/deliverables/기술표준서.docx','cpf-docs/deliverables/데이터베이스표준서.docx','cpf-docs/deliverables/산출물목록.docx']
PROV=re.compile(r'(Harness\s*(?:v)?\d+(?:\.\d+)+|Source\s*(?:SHA\s*)?[0-9A-F]{16,}|CPF_FULL_SOURCE_FOR_NEXT_QA_\d+)',re.I)
META=('누가 보는가','이 문서로 끝낼 일','기준')
NARROW=('번호','순번','id','상태','필수','기본값','yn','여부','유형','코드','code','type','status','default','required','no.')
WIDE=('설명','용도','사용','선택','조건','실패','복구','주의','대응','판단','경로','절차','결과','검증','api','option','옵션','reason','recovery','description','usage')
errors=[]
def fail(p,m): errors.append(f'{p}: {m}')
def text_of_table(t): return '\n'.join(c.text for r in t.rows for c in r.cells)
def para_has_nontext_object(p): return bool(p._p.xpath('.//w:drawing|.//w:pict|.//w:br[@w:type="page"]|.//w:fldChar|.//w:hyperlink'))
def grid_widths(t):
    try:return [int(x.get(qn('w:w'))) for x in t._tbl.tblGrid.findall(qn('w:gridCol'))]
    except:return []
def demand(s):
    s=(s or '').strip(); n=0.0
    for ch in s:
        o=ord(ch); n += 1.0 if (0xAC00<=o<=0xD7A3 or 0x4E00<=o<=0x9FFF) else (0.62 if ch.isalnum() else 0.35)
    return n
def col_demand(t,j):
    vals=[demand(r.cells[j].text) for r in t.rows if j<len(r.cells)]
    if not vals:return 0
    vals=sorted(vals); q=vals[min(len(vals)-1,max(0,math.ceil(len(vals)*0.75)-1))]
    return max(statistics.mean(vals),q*0.75)
def role(header):
    x=header.strip().lower()
    if any(k in x for k in NARROW): return 'narrow'
    if any(k in x for k in WIDE): return 'wide'
    return 'normal'
def symmetric(t):
    if not t.rows:return False
    hs=[c.text.strip().lower() for c in t.rows[0].cells]
    vendor=all(any(v in ' '.join(hs) for v in [x]) for x in ['oracle','postgresql','mariadb'])
    # also allow explicit same-role compare columns after one key column.
    if vendor:return True
    roles=[role(x) for x in hs]
    return len(set(roles))==1 and roles[0]=='normal'
for rel in DOCS:
    p=ROOT/rel
    if not p.is_file(): fail(rel,'missing'); continue
    try: doc=Document(p)
    except Exception as e: fail(rel,'open '+str(e)); continue
    # Vertical rhythm: numeric minimum + manual render gate required separately.
    for sn,min_before,min_after in [('Heading 1',52,11),('Heading 2',28,7),('Heading 3',18,5.5)]:
        if sn in doc.styles:
            pf=doc.styles[sn].paragraph_format
            b=pf.space_before.pt if pf.space_before else 0; a=pf.space_after.pt if pf.space_after else 0
            if b+0.01 < min_before or a+0.01 < min_after: fail(rel,f'{sn} vertical rhythm too tight {b}/{a}pt; min {min_before}/{min_after}')
    if 'Normal' in doc.styles:
        pf=doc.styles['Normal'].paragraph_format
        after=pf.space_after.pt if pf.space_after else 0
        if after+0.01<7.5: fail(rel,f'body paragraph spacing too tight {after}pt')
        try:
            ls=float(pf.line_spacing or 0)
            if ls and ls<1.25: fail(rel,f'body line spacing too tight {ls}')
        except: pass
    # Opening metadata table and provenance forbidden.
    for t in doc.tables[:2]:
        tt=text_of_table(t)
        if sum(1 for x in META if x in tt)>=2: fail(rel,'opening reader/purpose/basis metadata table')
    visible=[x.text for x in doc.paragraphs]
    for t in doc.tables: visible.append(text_of_table(t))
    for sec in doc.sections:
        visible.extend(x.text for x in sec.header.paragraphs); visible.extend(x.text for x in sec.footer.paragraphs)
    mt=PROV.search('\n'.join(visible))
    if mt: fail(rel,'user-facing production provenance '+mt.group(0))
    for idx,t in enumerate(doc.tables,1):
        if len(t.rows)==1: fail(rel,f'single-row layout table #{idx}')
        if not t.rows or not t.rows[0].cells: continue
        trPr=t.rows[0]._tr.get_or_add_trPr()
        if trPr.find(qn('w:tblHeader')) is None: fail(rel,f'table #{idx} header row not marked repeat header')
        gw=grid_widths(t); n=len(t.columns)
        if len(gw)!=n or n<2: continue
        total=sum(gw); shares=[w/total for w in gw]
        headers=[c.text.strip() for c in t.rows[0].cells]
        roles=[role(x) for x in headers]
        dem=[col_demand(t,j) for j in range(n)]
        mean=sum(dem)/n if n else 0
        variance=((max(dem)-min(dem))/mean*100) if mean else 0
        equal=(max(gw)-min(gw)) <= max(2,int(total/n*0.03))
        if equal and not symmetric(t) and (variance>12 or len(set(roles))>1): fail(rel,f'table #{idx} unjustified equal widths; demand variance={variance:.1f}% roles={roles}')
        # semantic width inversion: wide descriptive column cannot be narrower than narrow code/status column.
        if len(t.rows)>=3:
            for a,ra in enumerate(roles):
                for b,rb in enumerate(roles):
                    if ra=='wide' and rb=='narrow' and shares[a]+0.02<shares[b]:
                        fail(rel,f'table #{idx} semantic width inversion {headers[a]}({shares[a]:.1%}) < {headers[b]}({shares[b]:.1%})')
            for a in range(n):
                for b in range(n):
                    if dem[a]>=1.35*max(dem[b],1) and shares[a] < 1.15*shares[b] and roles[a] != 'narrow':
                        fail(rel,f'table #{idx} content-demand width mismatch col {a+1} demand {dem[a]:.1f} width {shares[a]:.1%} vs col {b+1}')
    # TOC/finder tab must fit writable width and show page number.
    if doc.sections:
        sec=doc.sections[0]; writable=sec.page_width-sec.left_margin-sec.right_margin
        for para in doc.paragraphs:
            if para.style and para.style.name=='CPF TOC Entry':
                tabs=list(para.paragraph_format.tab_stops)
                if not tabs: fail(rel,'TOC entry missing right tab stop: '+para.text[:50]); continue
                if max(t.position for t in tabs)>writable: fail(rel,'TOC tab stop outside writable width: '+para.text[:50])
                if not re.search(r'\t\s*\d+\s*$',para.text): fail(rel,'TOC page number missing/not visible: '+para.text[:50])
    # semantic heading styles and blank spacer paragraphs.
    prev=0
    for para in doc.paragraphs:
        st=para.style.name if para.style else ''
        if st.startswith('Heading '):
            try:lvl=int(st.split()[-1])
            except:continue
            if prev and lvl>prev+1: fail(rel,f'heading level skip {prev}->{lvl}: {para.text[:50]}')
            prev=lvl
    run=0
    for para in doc.paragraphs:
        if para.text.strip()=='' and not para_has_nontext_object(para): run+=1
        else: run=0
        if run>=2: fail(rel,'two or more consecutive blank spacer paragraphs'); break
    try:
        with zipfile.ZipFile(p) as z:
            if 'word/document.xml' not in z.namelist(): fail(rel,'word/document.xml missing')
    except Exception as e: fail(rel,'invalid package '+str(e))
if errors:
    print('DOCX_STRUCTURE=FAIL COUNT='+str(len(errors)))
    for e in errors: print('-',e)
    sys.exit(1)
print('DOCX_STRUCTURE=PASS DOCX='+str(len(DOCS)))
print('NOTE=Rendered wrap/vertical-rhythm evidence is still mandatory; OOXML pass alone is not final PASS.')
