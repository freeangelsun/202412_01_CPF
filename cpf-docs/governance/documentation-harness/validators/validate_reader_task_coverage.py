#!/usr/bin/env python3
import json,sys,re
from pathlib import Path
from docx import Document
H=Path(__file__).resolve().parents[1]
ROOT=H.parents[2]
C=json.loads((H/'reader-task-coverage.json').read_text(encoding='utf-8'))
errs=[]
def visible(path):
    if path.suffix.lower()=='.md': return path.read_text(encoding='utf-8')
    d=Document(str(path)); out=[p.text for p in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            out += [c.text for c in r.cells]
    return '\n'.join(out)
for a in C['artifacts']:
    p=ROOT/a['path']
    if not p.is_file(): errs.append(f"{a['id']}: missing {a['path']}"); continue
    txt=visible(p).lower()
    for concept in a['allRequiredConcepts']:
        if concept.lower() not in txt: errs.append(f"{a['id']}: missing concept {concept}")
if errs:
    print('READER_TASK_COVERAGE=FAIL COUNT='+str(len(errs)))
    for e in errs: print('-',e)
    sys.exit(1)
print('READER_TASK_COVERAGE=PASS ARTIFACTS='+str(len(C['artifacts'])))
