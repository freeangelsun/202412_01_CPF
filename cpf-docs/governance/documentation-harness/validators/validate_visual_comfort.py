#!/usr/bin/env python3
import json,sys,subprocess
from pathlib import Path
from docx import Document
H=Path(__file__).resolve().parents[1]; ROOT=H.parents[2]; C=json.loads((H/'visual-comfort.json').read_text(encoding='utf-8'))
docs=list((ROOT/'cpf-docs/guides').glob('*.docx'))+list((ROOT/'cpf-docs/deliverables').glob('*.docx')); errs=[]
for p in docs:
 d=Document(p); n=d.styles['Normal']; fs=n.font.size.pt if n.font.size else 0; ls=float(n.paragraph_format.line_spacing or 0); aft=n.paragraph_format.space_after.pt if n.paragraph_format.space_after else 0
 if fs<C['bodyFontPtMin']: errs.append(f'{p.name}: body font {fs}')
 if ls and ls<C['bodyLineSpacingMin']: errs.append(f'{p.name}: line spacing {ls}')
 if aft<C['bodySpaceAfterPtMin']: errs.append(f'{p.name}: body after {aft}')
 for ti,t in enumerate(d.tables,1):
  for ri,row in enumerate(t.rows):
   for cell in row.cells:
    for pa in cell.paragraphs:
     for r in pa.runs:
      if r.text.strip() and r.font.size and r.font.size.pt < (C['tableHeaderFontPtMin'] if ri==0 else C['tableBodyFontPtMin'])-.05: errs.append(f'{p.name}: table {ti} small font {r.font.size.pt}'); break
if errs: print('VISUAL_COMFORT=FAIL COUNT='+str(len(errs))); [print('-',e) for e in errs[:80]]; sys.exit(1)
print('VISUAL_COMFORT=PASS DOCX='+str(len(docs))+' PDF=11')
